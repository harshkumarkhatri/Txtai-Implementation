# Take 1
# import txtai

# embeddings = txtai.Embeddings()
# print("line 4")
# embeddings.index(["Correct", "Not what we hoped"])
# print("line 6")
# print(embeddings.search("positive", 1))
# print("line 8")

# Take 2
# import os
# from pathlib import Path
# import PyPDF2
# import docx
# from collections import Counter
# import re
# import string

# class SimpleDocumentAnalyzer:
#     def __init__(self, folder_path):
#         self.folder_path = folder_path
#         self.documents = {}
#         self.processed_texts = {}

#     def read_pdf(self, file_path):
#         """Read content from PDF file"""
#         text = ""
#         try:
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text()
#         except Exception as e:
#             print(f"Error reading PDF {file_path}: {e}")
#         return text

#     def read_docx(self, file_path):
#         """Read content from DOCX file"""
#         text = ""
#         try:
#             doc = docx.Document(file_path)
#             text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
#         except Exception as e:
#             print(f"Error reading DOCX {file_path}: {e}")
#         return text

#     def read_txt(self, file_path):
#         """Read content from TXT file"""
#         try:
#             with open(file_path, 'r', encoding='utf-8') as file:
#                 return file.read()
#         except Exception as e:
#             try:
#                 # Try with a different encoding if UTF-8 fails
#                 with open(file_path, 'r', encoding='latin-1') as file:
#                     return file.read()
#             except Exception as e:
#                 print(f"Error reading TXT {file_path}: {e}")
#                 return ""

#     def load_documents(self):
#         """Load all documents from the specified folder"""
#         for file_path in Path(self.folder_path).glob('*'):
#             if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
#                 try:
#                     if file_path.suffix.lower() == '.pdf':
#                         content = self.read_pdf(file_path)
#                     elif file_path.suffix.lower() == '.docx':
#                         content = self.read_docx(file_path)
#                     else:  # .txt
#                         content = self.read_txt(file_path)
                    
#                     if content:
#                         self.documents[file_path.name] = content
#                         # Process the text
#                         self.processed_texts[file_path.name] = self.preprocess_text(content)
#                 except Exception as e:
#                     print(f"Error processing {file_path}: {e}")

#     def preprocess_text(self, text):
#         """Basic text preprocessing"""
#         # Convert to lowercase
#         text = text.lower()
#         # Remove punctuation
#         text = text.translate(str.maketrans('', '', string.punctuation))
#         # Remove extra whitespace
#         text = ' '.join(text.split())
#         return text

#     def get_word_frequency(self, text, n=10):
#         """Get most frequent words in text"""
#         words = text.split()
#         # Remove common English stop words
#         stop_words = {'the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'i', 
#                      'it', 'for', 'not', 'on', 'with', 'he', 'as', 'you', 'do', 'at'}
#         words = [word for word in words if word not in stop_words]
#         return Counter(words).most_common(n)

#     def find_dates(self, text):
#         """Find dates in text using regex"""
#         # This is a simple date pattern - can be expanded based on needs
#         date_pattern = r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}[/-]\d{1,2}[/-]\d{1,2}'
#         return re.findall(date_pattern, text)

#     def find_emails(self, text):
#         """Find email addresses in text"""
#         email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
#         return re.findall(email_pattern, text)

#     def get_document_summary(self, doc_name):
#         """Generate a basic summary of the document"""
#         if doc_name not in self.documents:
#             return "Document not found"
            
#         raw_text = self.documents[doc_name]
#         processed_text = self.processed_texts[doc_name]
        
#         # Get basic statistics
#         word_count = len(processed_text.split())
#         sentence_count = len(re.split(r'[.!?]+', raw_text))
        
#         # Get most frequent words
#         frequent_words = self.get_word_frequency(processed_text)
        
#         # Find dates and emails
#         dates = self.find_dates(raw_text)
#         emails = self.find_emails(raw_text)
        
#         return {
#             'word_count': word_count,
#             'sentence_count': sentence_count,
#             'most_frequent_words': frequent_words,
#             'dates_found': dates,
#             'emails_found': emails
#         }

#     def search_keyword(self, keyword, doc_name=None):
#         """Search for keyword in documents"""
#         results = {}
#         docs_to_search = [doc_name] if doc_name else self.documents.keys()
        
#         for doc in docs_to_search:
#             if doc in self.processed_texts:
#                 keyword = keyword.lower()
#                 text = self.processed_texts[doc]
#                 count = text.count(keyword)
#                 # Get context around keyword
#                 contexts = []
#                 words = text.split()
#                 for i, word in enumerate(words):
#                     if word == keyword:
#                         start = max(0, i - 3)
#                         end = min(len(words), i + 4)
#                         context = ' '.join(words[start:end])
#                         contexts.append(context)
#                 results[doc] = {
#                     'count': count,
#                     'contexts': contexts[:5]  # Limit to 5 contexts
#                 }
        
#         return results

# def main():
#     # Example usage
#     folder_path = "gtet_policies"
#     analyzer = SimpleDocumentAnalyzer(folder_path)
#     analyzer.load_documents()
    
#     # Get summary of all documents
#     for doc_name in analyzer.documents:
#         print(f"\nAnalyzing {doc_name}:")
#         summary = analyzer.get_document_summary(doc_name)
#         print(f"Word count: {summary['word_count']}")
#         print(f"Sentence count: {summary['sentence_count']}")
#         print("Most frequent words:")
#         for word, count in summary['most_frequent_words']:
#             print(f"  {word}: {count}")
#         if summary['dates_found']:
#             print(f"Dates found: {summary['dates_found']}")
#         if summary['emails_found']:
#             print(f"Emails found: {summary['emails_found']}")

#     # Example keyword search
#     keyword = "example"
#     print(f"\nSearching for '{keyword}':")
#     results = analyzer.search_keyword(keyword)
#     for doc, result in results.items():
#         if result['count'] > 0:
#             print(f"\nIn {doc}:")
#             print(f"Found {result['count']} occurrences")
#             print("Sample contexts:")
#             for context in result['contexts']:
#                 print(f"  ...{context}...")

# if __name__ == "__main__":
#     main()

# Take 3
import os
from pathlib import Path
import PyPDF2
import docx
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.tokenize import sent_tokenize

class DocumentQA:
    def __init__(self, folder_path):
        self.folder_path = folder_path
        self.documents = {}
        self.sentences = {}
        self.embeddings = {}
        
        # Download required NLTK data
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
        
        # Load the sentence transformer model
        print("Loading the model... This might take a moment.")
        self.model = SentenceTransformer('all-MiniLM-L6-v2')  # Lightweight model
        print("Model loaded successfully!")

    def read_pdf(self, file_path):
        """Read content from PDF file"""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text()
        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")
        return text

    def read_docx(self, file_path):
        """Read content from DOCX file"""
        text = ""
        try:
            doc = docx.Document(file_path)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")
        return text

    def read_txt(self, file_path):
        """Read content from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                print(f"Error reading TXT {file_path}: {e}")
                return ""

    def load_documents(self):
        """Load and process all documents from the specified folder"""
        print("Loading documents...")
        for file_path in Path(self.folder_path).glob('*'):
            if file_path.suffix.lower() in ['.pdf', '.docx', '.txt']:
                try:
                    # Read document based on file type
                    if file_path.suffix.lower() == '.pdf':
                        content = self.read_pdf(file_path)
                    elif file_path.suffix.lower() == '.docx':
                        content = self.read_docx(file_path)
                    else:  # .txt
                        content = self.read_txt(file_path)
                    
                    if content:
                        doc_name = file_path.name
                        self.documents[doc_name] = content
                        
                        # Split document into sentences
                        sentences = sent_tokenize(content)
                        self.sentences[doc_name] = sentences
                        
                        # Create embeddings for sentences
                        print(f"Creating embeddings for {doc_name}...")
                        embeddings = self.model.encode(sentences)
                        self.embeddings[doc_name] = embeddings
                        
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")
        
        print("Document loading and processing complete!")

    def find_most_relevant_sentences(self, question, n=3):
        """Find the most relevant sentences for a given question"""
        question_embedding = self.model.encode([question])[0]
        
        all_relevant = []
        for doc_name, doc_embeddings in self.embeddings.items():
            # Calculate similarities between question and all sentences
            similarities = cosine_similarity([question_embedding], doc_embeddings)[0]
            
            # Get top matches for this document
            top_indices = np.argsort(similarities)[-n:][::-1]
            
            for idx in top_indices:
                score = similarities[idx]
                if score > 0.3:  # Similarity threshold
                    all_relevant.append({
                        'document': doc_name,
                        'sentence': self.sentences[doc_name][idx],
                        'score': score
                    })
        
        # Sort all results by score
        all_relevant.sort(key=lambda x: x['score'], reverse=True)
        return all_relevant[:n]

    def answer_question(self, question):
        """Answer a question based on the document content"""
        relevant_sentences = self.find_most_relevant_sentences(question)
        
        if not relevant_sentences:
            return "I couldn't find any relevant information to answer your question."
        
        # Format the answer
        answer = "Based on the documents:\n\n"
        for i, rel in enumerate(relevant_sentences, 1):
            answer += f"{i}. From '{rel['document']}':\n"
            answer += f"   {rel['sentence']}\n"
            answer += f"   (Confidence: {rel['score']:.2f})\n\n"
        
        return answer

def main():
    # Example usage
    folder_path = "gtet_policies"
    qa_system = DocumentQA(folder_path)
    qa_system.load_documents()
    
    print("\nDocument QA System is ready!")
    print("Type 'quit' to exit")
    
    while True:
        question = input("\nWhat would you like to know about the documents? ").strip()
        
        if question.lower() == 'quit':
            break
        
        if not question:
            continue
        
        answer = qa_system.answer_question(question)
        print("\nAnswer:", answer)

if __name__ == "__main__":
    main()